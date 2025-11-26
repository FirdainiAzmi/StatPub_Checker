import streamlit as st
import re
from spellchecker import SpellChecker
import PyPDF2
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======================
# LOAD WORDLIST
# ======================
@st.cache_resource
def load_wordlists():
    with open("wordlists/id_wordlist.txt", "r", encoding="utf-8") as f:
        indo_words = set(f.read().splitlines())

    with open("wordlists/eng_wordlist.txt", "r", encoding="utf-8") as f:
        eng_words = set(f.read().splitlines())

    return indo_words, eng_words


indo_dict, eng_dict = load_wordlists()
spell_en = SpellChecker(language='en')


# ======================
# FILE READER
# ======================
def read_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for p in reader.pages:
        if p.extract_text():
            text += p.extract_text() + " "
    return text


def read_docx(file):
    doc = Document(file)
    text = " ".join([p.text for p in doc.paragraphs])
    return text


# ======================
# DETEKSI TYPO
# ======================
def detect_typo(text):
    words = re.findall(r'\b[a-zA-Z]+\b', text)
    typos = set()

    for w in words:
        wl = w.lower()

        if wl in indo_dict:  # kata Indonesia benar
            continue

        if wl in eng_dict:  # kata English benar
            continue

        if wl in spell_en:  # dicek SpellChecker EN
            continue

        typos.add(w)

    return sorted(list(typos))


# ======================
# DETEKSI KATA INGGRIS (italic)
# ======================
def detect_english_words(text):
    words = re.findall(r'\b[a-zA-Z]+\b', text)
    english_found = []

    for w in words:
        wl = w.lower()
        if wl in eng_dict or wl in spell_en:
            english_found.append(w)

    return sorted(set(english_found))


# ======================
# DETEKSI FORMAT PERSENTASE BENAR
# ======================
def detect_wrong_percentage(text):
    wrong = []

    # Format benar: 12.34%
    pattern_correct = r'\b\d+\.\d+%\b'
    pattern_any = r'\b\d+[%]\b|\b\d+[,]\d+[%]\b'

    all_found = re.findall(pattern_any, text)
    correct = re.findall(pattern_correct, text)

    for item in all_found:
        if item not in correct:
            wrong.append(item)

    return sorted(set(wrong))


# ======================
# CREATE HIGHLIGHTED DOCUMENT
# ======================
def add_highlight(run, color_word):
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color_word)  # red, yellow, cyan
    run._r.get_or_add_rPr().append(highlight)


def create_highlighted_doc(original_text, typos, eng_words, wrong_percent):
    doc = Document()

    # Pisahkan per baris/paragraf
    paragraphs = original_text.split("\n")

    for para_text in paragraphs:
        para = doc.add_paragraph()

        # Pisahkan kata dengan mempertahankan spasi original
        tokens = re.findall(r'\S+|\s+', para_text)

        for token in tokens:
            if token.isspace():
                # tambahkan spasi apa adanya
                para.add_run(token)
                continue

            clean = re.sub(r"[^\w%.,-]", "", token)
            run = para.add_run(token)

            # Highlight kategori tertentu
            if clean in typos:
                add_highlight(run, "red")
            elif clean in eng_words:
                add_highlight(run, "yellow")
            elif clean in wrong_percent:
                add_highlight(run, "cyan")

    return doc


# ======================
# STREAMLIT UI
# ======================
st.title("📘 Checker Publikasi BPS – Typo, English & Persentase")
st.write("Unggah PDF atau Word untuk dianalisis.")

uploaded = st.file_uploader("Upload file", type=["pdf", "docx"])

if uploaded:
    if uploaded.type == "application/pdf":
        text = read_pdf(uploaded)
    else:
        text = read_docx(uploaded)

    st.subheader("📄 Hasil Analisis")

    # TYPO
    typos = detect_typo(text)
    st.write("### ❌ Kemungkinan Typo:")
    if typos:
        st.error(typos)
    else:
        st.success("Tidak ditemukan typo.")

    # English words
    eng_words = detect_english_words(text)
    st.write("### 🔤 Kata Bahasa Inggris (seharusnya italic):")
    if eng_words:
        st.warning(eng_words)
    else:
        st.success("Tidak ada kata Bahasa Inggris.")

    # Percentage wrong
    wrong_percent = detect_wrong_percentage(text)
    st.write("### % Format Persentase Salah:")
    if wrong_percent:
        st.warning(wrong_percent)
    else:
        st.success("Format persentase sudah benar (contoh benar: 12.45%).")

    # ======================
    # DOWNLOAD DOCUMENT
    # ======================
    st.write("---")
    st.write("### 📄 Download Dokumen Hasil Highlight")

    if st.button("Generate DOCX"):
        highlighted_doc = create_highlighted_doc(text, typos, eng_words, wrong_percent)
        output_path = "hasil_cek_publikasi.docx"
        highlighted_doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button(
                label="⬇ Download Hasil (DOCX)",
                data=f,
                file_name="hasil_cek_publikasi.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

